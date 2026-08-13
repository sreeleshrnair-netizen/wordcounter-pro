from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware

from pathlib import Path
from io import BytesIO

from docx import Document

import fitz
import csv
import re
import os


# =========================================================
# APPLICATION
# =========================================================

app = FastAPI(
    title="WordCounter Pro",
    version="2.0"
)


# =========================================================
# CONFIGURATION
# =========================================================

MAX_FILE_SIZE = 20 * 1024 * 1024

SUPPORTED_FILES = {
    ".pdf",
    ".docx",
    ".txt",
    ".csv"
}


# =========================================================
# CORS
# =========================================================

frontend_origin = os.getenv(
    "FRONTEND_ORIGIN",
    "*"
)

if frontend_origin == "*":

    allowed_origins = ["*"]

else:

    allowed_origins = [
        origin.strip()
        for origin in frontend_origin.split(",")
        if origin.strip()
    ]


app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


# =========================================================
# WORD ANALYSIS ENGINE
# =========================================================

def analyze_text(text: str):

    # -----------------------------------------------------
    # Normalize line endings
    # -----------------------------------------------------

    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")


    # -----------------------------------------------------
    # Unicode-aware word detection
    # -----------------------------------------------------

    words = re.findall(
        r"[^\W_]+(?:['’\-][^\W_]+)*",
        text,
        flags=re.UNICODE
    )


    # -----------------------------------------------------
    # Trimmed text
    # -----------------------------------------------------

    trimmed_text = text.strip()


    # -----------------------------------------------------
    # Paragraphs
    # -----------------------------------------------------

    paragraphs = []

    if trimmed_text:

        paragraphs = [
            paragraph.strip()
            for paragraph in re.split(
                r"\n\s*\n+",
                trimmed_text
            )
            if paragraph.strip()
        ]


    # -----------------------------------------------------
    # Sentences
    # -----------------------------------------------------

    sentence_matches = re.findall(
        r"[^.!?。！？]+[.!?。！？]+",
        trimmed_text
    )


    if trimmed_text:

        if sentence_matches:

            sentence_count = len(
                sentence_matches
            )

        else:

            # Text exists but contains no punctuation.
            sentence_count = 1

    else:

        sentence_count = 0


    # -----------------------------------------------------
    # Lines
    # -----------------------------------------------------

    if text:

        lines = [
            line
            for line in text.splitlines()
            if line.strip()
        ]

    else:

        lines = []


    # -----------------------------------------------------
    # Characters
    # -----------------------------------------------------

    characters = len(text)

    characters_without_spaces = len(
        re.sub(
            r"\s",
            "",
            text
        )
    )


    # -----------------------------------------------------
    # Word count
    # -----------------------------------------------------

    word_count = len(words)


    # -----------------------------------------------------
    # Reading time
    # Average reading speed = 200 WPM
    # -----------------------------------------------------

    if word_count > 0:

        reading_minutes = max(
            1,
            round(
                word_count / 200
            )
        )

    else:

        reading_minutes = 0


    # -----------------------------------------------------
    # Speaking time
    # Average speaking speed = 130 WPM
    # -----------------------------------------------------

    if word_count > 0:

        speaking_minutes = max(
            1,
            round(
                word_count / 130
            )
        )

    else:

        speaking_minutes = 0


    # -----------------------------------------------------
    # Average word length
    # -----------------------------------------------------

    if word_count > 0:

        clean_words = [
            re.sub(
                r"[^\w]",
                "",
                word,
                flags=re.UNICODE
            )
            for word in words
        ]

        total_length = sum(
            len(word)
            for word in clean_words
        )

        average_word_length = round(
            total_length / word_count,
            2
        )

    else:

        average_word_length = 0


    # -----------------------------------------------------
    # Normalize words
    # -----------------------------------------------------

    normalized_words = [
        word.casefold()
        for word in words
    ]


    # -----------------------------------------------------
    # Unique words
    # -----------------------------------------------------

    unique_words = set(
        normalized_words
    )


    # -----------------------------------------------------
    # Word frequency
    # -----------------------------------------------------

    frequency = {}

    for word in normalized_words:

        frequency[word] = (
            frequency.get(word, 0) + 1
        )


    # -----------------------------------------------------
    # Top 10 words
    # -----------------------------------------------------

    top_words = sorted(
        frequency.items(),
        key=lambda item: (
            -item[1],
            item[0]
        )
    )[:10]


    # -----------------------------------------------------
    # Longest word
    # -----------------------------------------------------

    longest_word = ""

    if words:

        longest_word = max(
            words,
            key=lambda word: len(
                re.sub(
                    r"[^\w]",
                    "",
                    word,
                    flags=re.UNICODE
                )
            )
        )


    # -----------------------------------------------------
    # Return analysis
    # -----------------------------------------------------

    return {

        "words": word_count,

        "characters": characters,

        "characters_no_spaces":
            characters_without_spaces,

        "paragraphs":
            len(paragraphs),

        "sentences":
            sentence_count,

        "lines":
            len(lines),

        "unique_words":
            len(unique_words),

        "reading_minutes":
            reading_minutes,

        "speaking_minutes":
            speaking_minutes,

        "average_word_length":
            average_word_length,

        "longest_word":
            longest_word,

        "top_words": [
            {
                "word": word,
                "count": count
            }
            for word, count in top_words
        ],

        "text": text
    }


# =========================================================
# PDF EXTRACTION
# =========================================================

def extract_pdf(data: bytes):

    try:

        document = fitz.open(
            stream=data,
            filetype="pdf"
        )

        pages = []

        for page in document:

            page_text = page.get_text(
                "text",
                sort=True
            )

            pages.append(
                page_text
            )

        page_count = len(
            document
        )

        document.close()

        text = "\n\n".join(
            pages
        )

        return text, page_count

    except Exception as error:

        raise ValueError(
            f"Unable to read PDF: {error}"
        )


# =========================================================
# DOCX EXTRACTION
# =========================================================

def extract_docx(data: bytes):

    try:

        document = Document(
            BytesIO(data)
        )

        content = []


        # -------------------------------------------------
        # Normal paragraphs
        # -------------------------------------------------

        for paragraph in document.paragraphs:

            if paragraph.text.strip():

                content.append(
                    paragraph.text
                )


        # -------------------------------------------------
        # Tables
        # -------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                row_text = []

                for cell in row.cells:

                    cell_text = (
                        cell.text.strip()
                    )

                    if cell_text:

                        row_text.append(
                            cell_text
                        )

                if row_text:

                    content.append(
                        " ".join(row_text)
                    )


        # -------------------------------------------------
        # Headers
        # -------------------------------------------------

        for section in document.sections:

            for paragraph in (
                section.header.paragraphs
            ):

                if paragraph.text.strip():

                    content.append(
                        paragraph.text
                    )


        # -------------------------------------------------
        # Footers
        # -------------------------------------------------

        for section in document.sections:

            for paragraph in (
                section.footer.paragraphs
            ):

                if paragraph.text.strip():

                    content.append(
                        paragraph.text
                    )


        return "\n".join(
            content
        ), None

    except Exception as error:

        raise ValueError(
            f"Unable to read DOCX: {error}"
        )


# =========================================================
# TXT EXTRACTION
# =========================================================

def extract_txt(data: bytes):

    return (
        data.decode(
            "utf-8-sig",
            errors="replace"
        ),
        None
    )


# =========================================================
# CSV EXTRACTION
# =========================================================

def extract_csv(data: bytes):

    decoded = data.decode(
        "utf-8-sig",
        errors="replace"
    )

    rows = csv.reader(
        decoded.splitlines()
    )

    text = []

    for row in rows:

        cleaned_row = [
            cell.strip()
            for cell in row
            if cell.strip()
        ]

        if cleaned_row:

            text.append(
                " ".join(
                    cleaned_row
                )
            )

    return (
        "\n".join(text),
        None
    )


# =========================================================
# HEALTH CHECK
# =========================================================

@app.get(
    "/api/health"
)
def health():

    return {
        "status": "online",
        "service": "WordCounter Pro",
        "version": "2.0"
    }


# =========================================================
# ROOT & STATUS
# =========================================================

@app.get(
    "/",
    include_in_schema=False
)
async def root():

    return {
        "message": "WordCounter Pro API is running.",
        "docs": "/docs"
    }


# =========================================================
# FILE ANALYSIS
# =========================================================

@app.post(
    "/api/analyze-file"
)
async def analyze_file(
    file: UploadFile = File(...)
):

    filename = (
        file.filename
        or "document"
    )


    # -----------------------------------------------------
    # File extension
    # -----------------------------------------------------

    extension = Path(
        filename
    ).suffix.lower()


    # -----------------------------------------------------
    # Validate file type
    # -----------------------------------------------------

    if extension not in SUPPORTED_FILES:

        raise HTTPException(
            status_code=400,
            detail=(
                "Unsupported file type. "
                "Supported formats: "
                "PDF, DOCX, TXT, CSV."
            )
        )


    # -----------------------------------------------------
    # Read file
    # -----------------------------------------------------

    data = await file.read()


    # -----------------------------------------------------
    # Validate file size
    # -----------------------------------------------------

    if len(data) > MAX_FILE_SIZE:

        raise HTTPException(
            status_code=413,
            detail=(
                "File exceeds the "
                "20 MB limit."
            )
        )


    # -----------------------------------------------------
    # Empty file
    # -----------------------------------------------------

    if len(data) == 0:

        raise HTTPException(
            status_code=400,
            detail="The uploaded file is empty."
        )


    # -----------------------------------------------------
    # Extract text
    # -----------------------------------------------------

    try:

        pages = None


        if extension == ".pdf":

            text, pages = extract_pdf(
                data
            )


        elif extension == ".docx":

            text, pages = extract_docx(
                data
            )


        elif extension == ".txt":

            text, pages = extract_txt(
                data
            )


        elif extension == ".csv":

            text, pages = extract_csv(
                data
            )


        else:

            text = ""


        # -------------------------------------------------
        # Analyze extracted text
        # -------------------------------------------------

        result = analyze_text(
            text
        )


        # -------------------------------------------------
        # File information
        # -------------------------------------------------

        result["filename"] = filename

        result["file_type"] = (
            extension
            .replace(
                ".",
                ""
            )
            .upper()
        )

        result["file_size"] = len(
            data
        )

        result["pages"] = pages


        # -------------------------------------------------
        # OCR detection
        # -------------------------------------------------

        if (
            extension == ".pdf"
            and not text.strip()
        ):

            result["ocr_required"] = True

        else:

            result["ocr_required"] = False


        return result


    except HTTPException:

        raise


    except Exception as error:

        raise HTTPException(
            status_code=422,
            detail=(
                f"Could not process file: "
                f"{error}"
            )
        )